"""Generate a Dossier de Competences (PDF / DOCX) from candidate parsed CV data."""

import io
from datetime import datetime, timezone

# ── PDF generation (reportlab) ──────────────────────────────────────────────
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import cm
from reportlab.platypus import (
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
    PageBreak,
)

BRAND = colors.HexColor("#2E86AB")
BRAND_LIGHT = colors.HexColor("#EDF6FA")
GRAY_800 = colors.HexColor("#1F2937")
GRAY_600 = colors.HexColor("#4B5563")
GRAY_200 = colors.HexColor("#E5E7EB")
WHITE = colors.white

PAGE_W, PAGE_H = A4


def _styles():
    ss = getSampleStyleSheet()
    return {
        "title": ParagraphStyle("dc_title", parent=ss["Title"], fontSize=22, textColor=BRAND, spaceAfter=6, fontName="Helvetica-Bold"),
        "subtitle": ParagraphStyle("dc_sub", parent=ss["Normal"], fontSize=13, textColor=GRAY_600, spaceAfter=18, fontName="Helvetica"),
        "h1": ParagraphStyle("dc_h1", parent=ss["Heading1"], fontSize=16, textColor=BRAND, spaceBefore=18, spaceAfter=10, fontName="Helvetica-Bold", borderWidth=0),
        "h2": ParagraphStyle("dc_h2", parent=ss["Heading2"], fontSize=12, textColor=GRAY_800, spaceBefore=12, spaceAfter=6, fontName="Helvetica-Bold"),
        "body": ParagraphStyle("dc_body", parent=ss["Normal"], fontSize=10, textColor=GRAY_800, leading=14, fontName="Helvetica"),
        "bullet": ParagraphStyle("dc_bullet", parent=ss["Normal"], fontSize=10, textColor=GRAY_800, leading=14, fontName="Helvetica", leftIndent=18, bulletIndent=6),
        "small": ParagraphStyle("dc_small", parent=ss["Normal"], fontSize=8, textColor=GRAY_600, fontName="Helvetica"),
    }


def _header_footer(canvas, doc, name: str):
    canvas.saveState()
    # Header line
    canvas.setStrokeColor(BRAND)
    canvas.setLineWidth(2)
    canvas.line(2 * cm, PAGE_H - 1.5 * cm, PAGE_W - 2 * cm, PAGE_H - 1.5 * cm)
    # Footer
    canvas.setFont("Helvetica", 8)
    canvas.setFillColor(GRAY_600)
    canvas.drawString(2 * cm, 1.2 * cm, f"Dossier de competences — {name}")
    canvas.drawRightString(PAGE_W - 2 * cm, 1.2 * cm, f"Page {doc.page}")
    canvas.restoreState()


def generate_dossier_pdf(data: dict) -> bytes:
    """Generate a professional PDF dossier de competences."""
    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4, leftMargin=2 * cm, rightMargin=2 * cm, topMargin=2 * cm, bottomMargin=2 * cm)
    s = _styles()
    story = []
    name = data.get("name", "Candidat")

    # ── Page 1: Header + Contact + Competences ──
    story.append(Paragraph("DOSSIER DE COMPETENCES", s["title"]))
    story.append(Spacer(1, 4))

    # Contact info
    contact_parts = []
    if data.get("email"):
        contact_parts.append(data["email"])
    if data.get("phone"):
        contact_parts.append(data["phone"])
    story.append(Paragraph(f"<b>{name}</b>", ParagraphStyle("name", parent=s["body"], fontSize=14, fontName="Helvetica-Bold")))
    if contact_parts:
        story.append(Paragraph(" | ".join(contact_parts), s["body"]))
    story.append(Spacer(1, 6))

    # Summary
    if data.get("summary"):
        story.append(Paragraph(data["summary"], s["body"]))
    story.append(Spacer(1, 12))

    # ── Competences techniques ──
    skills = data.get("skills", [])
    if skills:
        story.append(Paragraph("Domaines de competences", s["h1"]))
        for skill in skills:
            story.append(Paragraph(f"• {skill}", s["bullet"]))
        story.append(Spacer(1, 8))

    # ── Environnements techniques ──
    # Group skills by category if available
    tech_env = data.get("technical_environments")
    if tech_env and isinstance(tech_env, list):
        story.append(Paragraph("Environnements techniques", s["h1"]))
        for env in tech_env:
            story.append(Paragraph(f"◆ {env}", s["bullet"]))
        story.append(Spacer(1, 8))

    # ── Languages ──
    languages = data.get("languages", [])
    if languages:
        story.append(Paragraph("Langues", s["h1"]))
        for lang in languages:
            if isinstance(lang, dict):
                story.append(Paragraph(f"◆ <b>{lang.get('name', lang)}</b> — {lang.get('level', '')}", s["bullet"]))
            else:
                story.append(Paragraph(f"◆ {lang}", s["bullet"]))
        story.append(Spacer(1, 8))

    # ── Synthese des experiences (table) ──
    experiences = data.get("experiences", [])
    if experiences:
        story.append(Paragraph("Synthese des experiences", s["h1"]))
        table_data = [["Entreprise", "Poste", "Duree", "Competences cles"]]
        for exp in experiences:
            company = exp.get("company", "—")
            title = exp.get("title", "—")
            duration = exp.get("duration", "—")
            desc = exp.get("description", "")
            # Truncate description for table
            desc_short = desc[:80] + "..." if len(desc) > 80 else desc
            table_data.append([
                Paragraph(company, s["small"]),
                Paragraph(title, s["small"]),
                Paragraph(str(duration), s["small"]),
                Paragraph(desc_short, s["small"]),
            ])

        col_widths = [3.5 * cm, 3.5 * cm, 2.5 * cm, 7 * cm]
        t = Table(table_data, colWidths=col_widths, repeatRows=1)
        t.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), BRAND),
            ("TEXTCOLOR", (0, 0), (-1, 0), WHITE),
            ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
            ("FONTSIZE", (0, 0), (-1, 0), 9),
            ("FONTSIZE", (0, 1), (-1, -1), 8),
            ("ALIGN", (0, 0), (-1, 0), "CENTER"),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ("GRID", (0, 0), (-1, -1), 0.5, GRAY_200),
            ("ROWBACKGROUNDS", (0, 1), (-1, -1), [WHITE, BRAND_LIGHT]),
            ("TOPPADDING", (0, 0), (-1, -1), 6),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
            ("LEFTPADDING", (0, 0), (-1, -1), 6),
            ("RIGHTPADDING", (0, 0), (-1, -1), 6),
        ]))
        story.append(t)
        story.append(Spacer(1, 12))

    # ── Detail des experiences ──
    if experiences:
        story.append(PageBreak())
        story.append(Paragraph("Experiences detaillees", s["h1"]))

        for exp in experiences:
            company = exp.get("company", "")
            title = exp.get("title", "")
            duration = exp.get("duration", "")
            description = exp.get("description", "")

            story.append(Paragraph(f"<b>{company}</b>", ParagraphStyle("exp_company", parent=s["h2"], textColor=BRAND, fontSize=12)))
            meta_parts = []
            if title:
                meta_parts.append(f"<b>Poste :</b> {title}")
            if duration:
                meta_parts.append(f"<b>Duree :</b> {duration}")
            if meta_parts:
                story.append(Paragraph(" | ".join(meta_parts), s["body"]))

            if description:
                # Split description into bullet points if it contains newlines or is a list
                lines = [l.strip() for l in description.replace("•", "\n").split("\n") if l.strip()]
                if len(lines) > 1:
                    story.append(Paragraph("<b>Realisations :</b>", s["body"]))
                    for line in lines:
                        clean = line.lstrip("•-– ").strip()
                        if clean:
                            story.append(Paragraph(f"• {clean}", s["bullet"]))
                else:
                    story.append(Paragraph(description, s["body"]))

            story.append(Spacer(1, 10))

    # ── Formations ──
    education = data.get("education", [])
    if education:
        story.append(Paragraph("Formations et diplomes", s["h1"]))
        for edu in education:
            if isinstance(edu, dict):
                degree = edu.get("degree", "")
                school = edu.get("school", "")
                year = edu.get("year", "")
                story.append(Paragraph(f"◆ <b>{year}</b> — {degree} ({school})" if year else f"◆ {degree} — {school}", s["bullet"]))
            else:
                story.append(Paragraph(f"◆ {edu}", s["bullet"]))

    # ── Footer note ──
    story.append(Spacer(1, 20))
    story.append(Paragraph(
        f"Document genere le {datetime.now(timezone.utc).strftime('%d/%m/%Y')} par AIHM",
        ParagraphStyle("footer_note", parent=s["small"], alignment=1),
    ))

    doc.build(story, onFirstPage=lambda c, d: _header_footer(c, d, name), onLaterPages=lambda c, d: _header_footer(c, d, name))
    return buf.getvalue()


# ── DOCX generation (python-docx) ───────────────────────────────────────────

def generate_dossier_docx(data: dict) -> bytes:
    """Generate a professional DOCX dossier de competences."""
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT

    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    brand_color = RGBColor(0x2E, 0x86, 0xAB)
    dark_color = RGBColor(0x1F, 0x29, 0x37)
    gray_color = RGBColor(0x4B, 0x55, 0x63)

    name = data.get("name", "Candidat")

    # ── Title ──
    title = doc.add_heading("DOSSIER DE COMPETENCES", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in title.runs:
        run.font.color.rgb = brand_color
        run.font.size = Pt(22)

    # ── Contact ──
    p = doc.add_paragraph()
    run = p.add_run(name)
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = dark_color

    contact_parts = []
    if data.get("email"):
        contact_parts.append(data["email"])
    if data.get("phone"):
        contact_parts.append(data["phone"])
    if contact_parts:
        p2 = doc.add_paragraph(" | ".join(contact_parts))
        p2.runs[0].font.size = Pt(10)
        p2.runs[0].font.color.rgb = gray_color

    # Summary
    if data.get("summary"):
        p = doc.add_paragraph(data["summary"])
        p.runs[0].font.size = Pt(10)
        p.runs[0].font.color.rgb = dark_color
        p.paragraph_format.space_after = Pt(12)

    # ── Competences ──
    skills = data.get("skills", [])
    if skills:
        h = doc.add_heading("Domaines de competences", level=1)
        for run in h.runs:
            run.font.color.rgb = brand_color
        for skill in skills:
            p = doc.add_paragraph(skill, style="List Bullet")
            for run in p.runs:
                run.font.size = Pt(10)

    # ── Languages ──
    languages = data.get("languages", [])
    if languages:
        h = doc.add_heading("Langues", level=1)
        for run in h.runs:
            run.font.color.rgb = brand_color
        for lang in languages:
            text = f"{lang.get('name', lang)} — {lang.get('level', '')}" if isinstance(lang, dict) else str(lang)
            p = doc.add_paragraph(text, style="List Bullet")
            for run in p.runs:
                run.font.size = Pt(10)

    # ── Synthese des experiences (table) ──
    experiences = data.get("experiences", [])
    if experiences:
        h = doc.add_heading("Synthese des experiences", level=1)
        for run in h.runs:
            run.font.color.rgb = brand_color

        table = doc.add_table(rows=1, cols=4)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = "Table Grid"

        # Header
        headers = ["Entreprise", "Poste", "Duree", "Competences cles"]
        for i, text in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = text
            for p in cell.paragraphs:
                for run in p.runs:
                    run.bold = True
                    run.font.size = Pt(9)
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            from docx.oxml.ns import qn
            shading = cell._element.get_or_add_tcPr()
            shading_elem = shading.makeelement(qn("w:shd"), {qn("w:fill"): "2E86AB", qn("w:val"): "clear"})
            shading.append(shading_elem)

        for exp in experiences:
            row = table.add_row()
            row.cells[0].text = exp.get("company", "—")
            row.cells[1].text = exp.get("title", "—")
            row.cells[2].text = str(exp.get("duration", "—"))
            desc = exp.get("description", "")
            row.cells[3].text = desc[:80] + "..." if len(desc) > 80 else desc
            for cell in row.cells:
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.font.size = Pt(8)

    # ── Experiences detaillees ──
    if experiences:
        doc.add_page_break()
        h = doc.add_heading("Experiences detaillees", level=1)
        for run in h.runs:
            run.font.color.rgb = brand_color

        for exp in experiences:
            company = exp.get("company", "")
            title_text = exp.get("title", "")
            duration = exp.get("duration", "")
            description = exp.get("description", "")

            h2 = doc.add_heading(company, level=2)
            for run in h2.runs:
                run.font.color.rgb = brand_color
                run.font.size = Pt(12)

            if title_text or duration:
                p = doc.add_paragraph()
                if title_text:
                    run = p.add_run(f"Poste : {title_text}")
                    run.bold = True
                    run.font.size = Pt(10)
                if title_text and duration:
                    p.add_run(" | ")
                if duration:
                    run = p.add_run(f"Duree : {duration}")
                    run.font.size = Pt(10)

            if description:
                lines = [l.strip() for l in description.replace("•", "\n").split("\n") if l.strip()]
                if len(lines) > 1:
                    p = doc.add_paragraph()
                    run = p.add_run("Realisations :")
                    run.bold = True
                    run.font.size = Pt(10)
                    for line in lines:
                        clean = line.lstrip("•-– ").strip()
                        if clean:
                            bp = doc.add_paragraph(clean, style="List Bullet")
                            for r in bp.runs:
                                r.font.size = Pt(10)
                else:
                    p = doc.add_paragraph(description)
                    for r in p.runs:
                        r.font.size = Pt(10)

    # ── Formations ──
    education = data.get("education", [])
    if education:
        h = doc.add_heading("Formations et diplomes", level=1)
        for run in h.runs:
            run.font.color.rgb = brand_color
        for edu in education:
            if isinstance(edu, dict):
                degree = edu.get("degree", "")
                school = edu.get("school", "")
                year = edu.get("year", "")
                text = f"{year} — {degree} ({school})" if year else f"{degree} — {school}"
            else:
                text = str(edu)
            p = doc.add_paragraph(text, style="List Bullet")
            for r in p.runs:
                r.font.size = Pt(10)

    # ── Footer ──
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"\nDocument genere le {datetime.now(timezone.utc).strftime('%d/%m/%Y')} par AIHM")
    run.font.size = Pt(8)
    run.font.color.rgb = gray_color

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
